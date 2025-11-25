from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

NBSP = "\u00A0"  # неразрывный пробел


def set_font(run, name: str, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(14)
    run.font.bold = bold


def make_header_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_normal_paragraph(doc: Document, text: str):
    """Обычный абзац с красной строкой 1.25 см, TNR 14."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_class_heading(doc: Document, number: int, kind: str, name: str):
    """
    Заголовок: "N Класс Name:" или "N Структура Name:"
    N + 'Класс/Структура' — Times New Roman
    Name: — Courier New
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)

    r1 = p.add_run(f"{number} {kind} ")
    set_font(r1, "Times New Roman")

    r2 = p.add_run(f"{name}:")
    set_font(r2, "Courier New")

    return p


def add_bullet_block(doc: Document, title: str, items: list[dict]):
    """
    items: список словарей вида {"code": "...", "desc": "..."}.
    code — шрифт кода (Courier New), desc — Times New Roman.
    """
    # Заголовок "Поля:" / "Методы:"
    make_normal_paragraph(doc, title)

    if not items:
        return

    n = len(items)
    for idx, item in enumerate(items):
        last = (idx == n - 1)
        p = doc.add_paragraph()
        pf = p.paragraph_format

        # ВАЖНО: убираем красную строку принудительно
        pf.left_indent = Cm(2.5)
        pf.first_line_indent = Cm(0)   # ← отключаем красную строку
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)

        code = item.get("code", "").strip()
        desc = (item.get("desc", "") or "").strip()

        # Маркер "– " (с NBSP), Times New Roman
        r0 = p.add_run("–" + NBSP)
        set_font(r0, "Times New Roman")

        # Код — Courier New
        if code:
            r_code = p.add_run(code)
            set_font(r_code, "Courier New")

        # Описание — Times New Roman
        ending = "." if last else ";"

        if desc:
            r_colon = p.add_run(": ")
            set_font(r_colon, "Times New Roman")

            r_desc = p.add_run(desc + ending)
            set_font(r_desc, "Times New Roman")
        else:
            # Нет описания — просто код + точка/точка с запятой
            r_end = p.add_run(ending)
            set_font(r_end, "Times New Roman")



def build_sections() -> list[dict]:
    """
    Возвращает список всех 33 сущностей (классы/структуры) с полями/методами.
    Формат элемента:
    {
        "kind": "Класс" или "Структура",
        "name": "Имя",
        "desc": "Описание",
        "fields": [{"code": "...", "desc": "..."}, ...],
        "methods": [{"code": "...", "desc": "..."}, ...],
    }
    """
    sections: list[dict] = []

    # 1. Order
    sections.append({
        "kind": "Класс",
        "name": "Order",
        "desc": "Модель заказа: идентификатор, клиент, статус, позиции, сумма и момент создания.",
        "fields": [
            {"code": "int id", "desc": "идентификатор заказа"},
            {"code": "std::string client", "desc": "имя клиента"},
            {"code": "std::string status", "desc": "статус (new, in_progress, done, canceled)"},
            {"code": "std::map<std::string, int, std::less<>> items", "desc": "позиции заказа (товар → количество)"},
            {"code": "double total", "desc": "итоговая сумма заказа"},
            {"code": "std::string createdAt", "desc": "момент создания в формате ISO 8601"},
        ],
        "methods": [
            {
                "code": "double calcTotal(const std::map<std::string, double, std::less<>>& priceList) const",
                "desc": "расчёт суммы по прайс-листу"
            },
            {
                "code": "auto operator<=>(const Order& other) const",
                "desc": "сравнение заказов по идентификатору"
            },
            {
                "code": "bool operator==(const Order& other) const",
                "desc": "проверка равенства по идентификатору"
            },
            {
                "code": "std::string toLine() const",
                "desc": "сериализация заказа в CSV-подобную строку"
            },
            {
                "code": "static std::optional<Order> fromLine(const std::string& line)",
                "desc": "десериализация заказа из строки"
            },
            {
                "code": "friend std::ostream& operator<<(std::ostream& os, const Order& o)",
                "desc": "человекочитаемый вывод заказа в поток"
            },
        ],
    })

    # 2. Product
    sections.append({
        "kind": "Класс",
        "name": "Product",
        "desc": "Модель товара: наименование, цена и остаток на складе.",
        "fields": [
            {"code": "std::string name", "desc": "название товара"},
            {"code": "double price", "desc": "цена товара"},
            {"code": "int stock", "desc": "остаток товара на складе"},
        ],
        "methods": [
            {"code": "Product()", "desc": "конструктор по умолчанию (price = 0.0, stock = 0)"},
            {
                "code": "Product(const std::string& n, double p, int s)",
                "desc": "конструктор с параметрами (имя, цена, остаток)"
            },
        ],
    })

    # 3. FilterParams
    sections.append({
        "kind": "Структура",
        "name": "FilterParams",
        "desc": "Параметры предзаполнения диалога фильтрации заказов.",
        "fields": [
            {"code": "QString currentClient", "desc": "текущий фильтр по клиенту"},
            {"code": "QString currentStatus", "desc": "текущий фильтр по статусу"},
            {"code": "QString minTotal", "desc": "текущий нижний порог суммы"},
            {"code": "QString maxTotal", "desc": "текущий верхний порог суммы"},
            {"code": "QString minId", "desc": "минимальный ID"},
            {"code": "QString maxId", "desc": "максимальный ID"},
            {"code": "QDateTime fromDt", "desc": "начало диапазона дат"},
            {"code": "bool useFrom", "desc": "флаг использования нижней границы даты"},
            {"code": "QDateTime toDt", "desc": "конец диапазона дат"},
            {"code": "bool useTo", "desc": "флаг использования верхней границы даты"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения параметров"},
        ],
    })

    # 4. ReportFilterInfo
    sections.append({
        "kind": "Структура",
        "name": "ReportFilterInfo",
        "desc": "Параметры фильтрации, применённые к отчёту.",
        "fields": [
            {"code": "QString clientFilter", "desc": "применённый фильтр по клиенту"},
            {"code": "QString statusFilter", "desc": "применённый фильтр по статусу"},
            {"code": "QString minTotal", "desc": "нижняя граница суммы"},
            {"code": "QString maxTotal", "desc": "верхняя граница суммы"},
            {"code": "QString minId", "desc": "нижний ID"},
            {"code": "QString maxId", "desc": "верхний ID"},
            {"code": "QDateTime fromDate", "desc": "фактическая нижняя граница даты"},
            {"code": "QDateTime toDate", "desc": "фактическая верхняя граница даты"},
            {"code": "bool useFrom", "desc": "используется ли нижняя граница даты"},
            {"code": "bool useTo", "desc": "используется ли верхняя граница даты"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура-контейнер для информации о фильтре"},
        ],
    })

    # 5. OrderStatsLabels
    sections.append({
        "kind": "Структура",
        "name": "OrderStatsLabels",
        "desc": "Набор виджетов Qt для отображения статистики по заказам.",
        "fields": [
            {"code": "QLabel* newLabel_", "desc": "лейбл с количеством новых заказов"},
            {"code": "QLabel* inProgressLabel_", "desc": "лейбл с количеством заказов в работе"},
            {"code": "QLabel* doneLabel_", "desc": "лейбл с количеством выполненных заказов"},
            {"code": "QLabel* canceledLabel_", "desc": "лейбл с количеством отменённых заказов"},
            {"code": "QLabel* totalRevenueLabel_", "desc": "лейбл с общей выручкой"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "служебная структура для группировки лейблов"},
        ],
    })

    # 6. ProductStatsLabels
    sections.append({
        "kind": "Структура",
        "name": "ProductStatsLabels",
        "desc": "Набор лейблов для отображения агрегированной статистики по товарам.",
        "fields": [
            {"code": "QLabel* lowStockLabel_", "desc": "лейбл для товаров с низким остатком"},
            {"code": "QLabel* highStockLabel_", "desc": "лейбл для товаров с высоким остатком"},
            {"code": "QLabel* expensiveLabel_", "desc": "лейбл для самых дорогих товаров"},
            {"code": "QLabel* cheapLabel_", "desc": "лейбл для самых дешёвых товаров"},
            {"code": "QLabel* totalCountLabel_", "desc": "общее количество товаров"},
            {"code": "QLabel* totalValueLabel_", "desc": "общая стоимость запасов"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "служебная структура для группировки лейблов"},
        ],
    })

    # 7. FilterWidgets
    sections.append({
        "kind": "Структура",
        "name": "FilterWidgets",
        "desc": "Указатели на виджеты формы фильтрации в главном окне.",
        "fields": [
            {"code": "QLineEdit* clientEdit_", "desc": "поле ввода клиента"},
            {"code": "QComboBox* statusCombo_", "desc": "комбобокс выбора статуса"},
            {"code": "QLineEdit* minTotalEdit_", "desc": "поле нижней границы суммы"},
            {"code": "QLineEdit* maxTotalEdit_", "desc": "поле верхней границы суммы"},
            {"code": "QLineEdit* minIdEdit_", "desc": "поле нижней границы ID"},
            {"code": "QLineEdit* maxIdEdit_", "desc": "поле верхней границы ID"},
            {"code": "QCheckBox* useFromCheck_", "desc": "чекбокс использовать нижнюю границу даты"},
            {"code": "QDateTimeEdit* fromDateEdit_", "desc": "выбор нижней даты"},
            {"code": "QCheckBox* useToCheck_", "desc": "чекбокс использовать верхнюю границу даты"},
            {"code": "QDateTimeEdit* toDateEdit_", "desc": "выбор верхней даты"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура-контейнер для указателей на виджеты"},
        ],
    })

    # 8. FilterState
    sections.append({
        "kind": "Структура",
        "name": "FilterState",
        "desc": "Текущее состояние фильтров в главном окне.",
        "fields": [
            {"code": "QString activeClientFilter_", "desc": "активный фильтр по клиенту"},
            {"code": "QString activeStatusFilter_", "desc": "активный фильтр по статусу"},
            {"code": "QString minTotalText_", "desc": "текст нижней границы суммы"},
            {"code": "QString maxTotalText_", "desc": "текст верхней границы суммы"},
            {"code": "QString minIdText_", "desc": "текст нижней границы ID"},
            {"code": "QString maxIdText_", "desc": "текст верхней границы ID"},
            {"code": "QDateTime fromDate_", "desc": "активная нижняя граница дат"},
            {"code": "QDateTime toDate_", "desc": "активная верхняя граница дат"},
            {"code": "bool useFrom_", "desc": "флаг использования нижней границы даты"},
            {"code": "bool useTo_", "desc": "флаг использования верхней границы даты"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура для хранения состояния фильтра"},
        ],
    })

    # 9. StatusStats
    sections.append({
        "kind": "Структура",
        "name": "StatusStats",
        "desc": "Агрегированные данные по количеству и выручке по каждому статусу заказа.",
        "fields": [
            {"code": "int newCount", "desc": "количество заказов в статусе new"},
            {"code": "int inProgressCount", "desc": "количество заказов в статусе in_progress"},
            {"code": "int doneCount", "desc": "количество заказов в статусе done"},
            {"code": "int canceledCount", "desc": "количество заказов в статусе canceled"},
            {"code": "double newRevenue", "desc": "выручка по заказам new"},
            {"code": "double inProgressRevenue", "desc": "выручка по заказам in_progress"},
            {"code": "double doneRevenue", "desc": "выручка по заказам done"},
            {"code": "double canceledRevenue", "desc": "выручка по заказам canceled"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "служебная структура для окна статистики"},
        ],
    })

    # 10. RevenueBarParams
    sections.append({
        "kind": "Структура",
        "name": "RevenueBarParams",
        "desc": "Параметры отрисовки столбца выручки в окне статистики.",
        "fields": [
            {"code": "double revenue", "desc": "значение выручки"},
            {"code": "QColor color", "desc": "цвет столбца"},
            {"code": "QString statusName", "desc": "подпись статуса"},
            {"code": "int x", "desc": "координата X столбца"},
            {"code": "int baseY", "desc": "базовая координата Y (основание)"},
            {"code": "int barWidth", "desc": "ширина столбца"},
            {"code": "int chartHeight", "desc": "высота области диаграммы"},
            {"code": "double maxRevenue", "desc": "максимальная выручка для нормировки"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура параметров для отрисовки столбца"},
        ],
    })

    # 11. IRepository
    sections.append({
        "kind": "Класс",
        "name": "IRepository",
        "desc": "Интерфейс репозитория заказов.",
        "fields": [
            {"code": "virtual ~IRepository()", "desc": "виртуальный деструктор"},
        ],
        "methods": [
            {
                "code": "virtual void save(const std::vector<Order>& data) = 0",
                "desc": "сохранить в хранилище набор заказов"
            },
            {
                "code": "virtual std::vector<Order> load() = 0",
                "desc": "загрузить набор заказов из хранилища"
            },
        ],
    })

    # 12. IProductRepository
    sections.append({
        "kind": "Класс",
        "name": "IProductRepository",
        "desc": "Интерфейс репозитория товаров/прайс-листа.",
        "fields": [
            {"code": "virtual ~IProductRepository()", "desc": "виртуальный деструктор"},
        ],
        "methods": [
            {
                "code": "virtual void save(const std::map<std::string, Product, std::less<>>& data) = 0",
                "desc": "сохранить набор товаров"
            },
            {
                "code": "virtual std::map<std::string, Product, std::less<>> load() = 0",
                "desc": "загрузить набор товаров"
            },
        ],
    })

    # 13. TxtOrderRepository
    sections.append({
        "kind": "Класс",
        "name": "TxtOrderRepository",
        "desc": "Файловая реализация IRepository для хранения заказов в текстовом файле.",
        "fields": [
            {"code": "std::string file_", "desc": "путь к файлу с заказами"},
        ],
        "methods": [
            {
                "code": "explicit TxtOrderRepository(std::string f)",
                "desc": "конструктор с путём к файлу"
            },
            {
                "code": "void save(const std::vector<Order>& data) override",
                "desc": "записать заказы в файл"
            },
            {
                "code": "std::vector<Order> load() override",
                "desc": "прочитать заказы из файла"
            },
        ],
    })

    # 14. TxtProductRepository
    sections.append({
        "kind": "Класс",
        "name": "TxtProductRepository",
        "desc": "Файловая реализация IProductRepository для хранения прайс-листа в текстовом файле.",
        "fields": [
            {"code": "std::string file_", "desc": "путь к файлу прайс-листа"},
        ],
        "methods": [
            {
                "code": "explicit TxtProductRepository(std::string f)",
                "desc": "конструктор с путём к файлу"
            },
            {
                "code": "void save(const std::map<std::string, Product, std::less<>>& data) override",
                "desc": "сохранить набор товаров в файл"
            },
            {
                "code": "std::map<std::string, Product, std::less<>> load() override",
                "desc": "загрузить товары из файла"
            },
        ],
    })

    # 15. OrderService
    sections.append({
        "kind": "Класс",
        "name": "OrderService",
        "desc": "Сервис бизнес-логики заказов: создание, изменение, пересчёт и персистентность.",
        "fields": [
            {"code": "SimpleList_Order data_", "desc": "оперативное хранилище заказов"},
            {"code": "std::map<std::string, double, std::less<>> price_", "desc": "прайс-лист (товар → цена)"},
            {"code": "int nextId_", "desc": "счётчик следующего идентификатора заказа"},
            {"code": "IRepository& repo_", "desc": "репозиторий заказов"},
            {"code": "ProductService* productService_", "desc": "сервис товаров для учёта остатков"},
        ],
        "methods": [
            {"code": "explicit OrderService(IRepository& repo)", "desc": "конструктор сервиса заказов"},
            {"code": "void setProductService(ProductService* ps)", "desc": "установить сервис товаров"},
            {
                "code": "void setPrices(const std::map<std::string, Product, std::less<>>& products)",
                "desc": "обновить прайс-лист по коллекции товаров"
            },
            {
                "code": "const std::map<std::string, double, std::less<>>& price() const",
                "desc": "получить текущий прайс-лист"
            },
            {
                "code": "Order& create(const std::string& client)",
                "desc": "создать новый заказ для клиента"
            },
            {
                "code": "void addItem(Order& o, const std::string& name, int qty)",
                "desc": "добавить позицию в заказ с учётом остатков"
            },
            {
                "code": "void removeItem(Order& o, const std::string& name)",
                "desc": "удалить позицию из заказа с возвратом товара на склад"
            },
            {
                "code": "void setStatus(Order& o, const std::string& s)",
                "desc": "изменить статус заказа с учётом остатков"
            },
            {
                "code": "Order* findById(int id)",
                "desc": "найти заказ по идентификатору (неconst)"
            },
            {
                "code": "const Order* findById(int id) const",
                "desc": "найти заказ по идентификатору (const-вариант)"
            },
            {"code": "void sortById()", "desc": "отсортировать заказы по идентификатору"},
            {"code": "double revenue() const", "desc": "подсчитать суммарную выручку"},
            {
                "code": "void recalculateOrdersWithProduct(const std::string& productKey)",
                "desc": "пересчитать заказы, затронутые изменением товара"
            },
            {"code": "void save()", "desc": "сохранить заказы в репозиторий"},
            {"code": "void load()", "desc": "загрузить заказы из репозитория"},
            {
                "code": "const SimpleList_Order& all() const",
                "desc": "получить доступ ко всем заказам"
            },
            {
                "code": "int& nextIdRef()",
                "desc": "получить ссылку на счётчик идентификаторов"
            },
        ],
    })

    # 16. ProductService
    sections.append({
        "kind": "Класс",
        "name": "ProductService",
        "desc": "Сервис управления товарами и ценами (CRUD, валидация, загрузка/сохранение).",
        "fields": [
            {"code": "std::map<std::string, Product, std::less<>> products_", "desc": "коллекция товаров"},
            {"code": "IProductRepository& repo_", "desc": "репозиторий товаров"},
            {"code": "ValidationService V_", "desc": "сервис валидации"},
        ],
        "methods": [
            {"code": "explicit ProductService(IProductRepository& repo)", "desc": "конструктор сервиса товаров"},
            {
                "code": "const std::map<std::string, Product, std::less<>>& all() const",
                "desc": "доступ ко всем товарам"
            },
            {
                "code": "Product* findProduct(const std::string& name)",
                "desc": "поиск товара по имени (изменяемый указатель)"
            },
            {
                "code": "const Product* findProduct(const std::string& name) const",
                "desc": "поиск товара по имени (const-указатель)"
            },
            {"code": "void load()", "desc": "загрузить товары из репозитория"},
            {"code": "void save()", "desc": "сохранить товары в репозиторий"},
            {
                "code": "void addProduct(const std::string& name, double price, int stock)",
                "desc": "добавить новый товар"
            },
            {
                "code": "void removeProduct(const std::string& name)",
                "desc": "удалить товар по имени"
            },
            {
                "code": "void updateProduct(const std::string& oldName, const std::string& newName, double newPrice, int stock)",
                "desc": "обновить параметры товара"
            },
            {
                "code": "void decreaseStock(const std::string& name, int qty)",
                "desc": "уменьшить остаток товара"
            },
            {
                "code": "void increaseStock(const std::string& name, int qty)",
                "desc": "увеличить остаток товара"
            },
            {
                "code": "bool hasEnoughStock(const std::string& name, int qty) const",
                "desc": "проверить достаточность остатка"
            },
            {
                "code": "int getStock(const std::string& name) const",
                "desc": "получить текущий остаток"
            },
        ],
    })

    # 17. ReportService
    sections.append({
        "kind": "Класс",
        "name": "ReportService",
        "desc": "Сервис формирования CSV-отчётов по заказам.",
        "fields": [
            {"code": "(нет)", "desc": "логика реализована статическими методами"},
        ],
        "methods": [
            {
                "code": "static QString generateReport(const QList<const Order*>& orders, const QString& reportName, bool scopeFiltered, bool includeFiltersHeader, bool includeSummarySection, const ReportFilterInfo& filterInfo, const OrderService& orderService)",
                "desc": "сформировать текст отчёта по заказам"
            },
            {
                "code": "static QString sanitizedBaseName(const QString& raw)",
                "desc": "нормализовать имя файла отчёта"
            },
            {
                "code": "static QString escapeCsvField(const QString& field)",
                "desc": "экранировать поле для CSV"
            },
        ],
    })

    # 18. CustomException
    sections.append({
        "kind": "Класс",
        "name": "CustomException",
        "desc": "Базовый класс прикладных исключений (наследуется от std::runtime_error).",
        "fields": [
            {"code": "(нет)", "desc": "дополнительных полей не содержит"},
        ],
        "methods": [
            {
                "code": "explicit CustomException(const std::string& message)",
                "desc": "создать исключение с сообщением"
            },
        ],
    })

    # 19. ValidationException
    sections.append({
        "kind": "Класс",
        "name": "ValidationException",
        "desc": "Исключение ошибок валидации входных данных.",
        "fields": [
            {"code": "(нет)", "desc": "наследует поведение CustomException"},
        ],
        "methods": [
            {
                "code": "explicit ValidationException(const std::string& message)",
                "desc": "создать исключение валидации"
            },
        ],
    })

    # 20. NotFoundException
    sections.append({
        "kind": "Класс",
        "name": "NotFoundException",
        "desc": "Исключение, сигнализирующее об отсутствии сущности.",
        "fields": [
            {"code": "(нет)", "desc": "наследует поведение CustomException"},
        ],
        "methods": [
            {
                "code": "explicit NotFoundException(const std::string& message)",
                "desc": "создать исключение \"не найдено\""
            },
        ],
    })

    # 21. IoException
    sections.append({
        "kind": "Класс",
        "name": "IoException",
        "desc": "Исключение ошибок ввода-вывода.",
        "fields": [
            {"code": "(нет)", "desc": "наследует поведение CustomException"},
        ],
        "methods": [
            {
                "code": "explicit IoException(const std::string& message)",
                "desc": "создать исключение ввода-вывода"
            },
        ],
    })

    # 22. ValidationService
    sections.append({
        "kind": "Класс",
        "name": "ValidationService",
        "desc": "Сервис валидации строк, чисел и денежных значений.",
        "fields": [
            {"code": "std::regex re", "desc": "внутренний шаблон для проверок"},
        ],
        "methods": [
            {
                "code": "void validate_client_name(const std::string& name)",
                "desc": "проверить корректность имени клиента"
            },
            {
                "code": "void validate_item_name(const std::string& name)",
                "desc": "проверить корректность названия товара"
            },
            {
                "code": "void validate_qty(int qty)",
                "desc": "проверить количество (> 0)"
            },
            {
                "code": "void validate_status(std::string_view s)",
                "desc": "проверить корректность статуса заказа"
            },
            {
                "code": "void validate_id(int id)",
                "desc": "проверить положительный идентификатор"
            },
            {
                "code": "void validate_price(double price)",
                "desc": "проверить корректность цены (положительная, 2 знака)"
            },
            {
                "code": "double normalize_money(double v)",
                "desc": "нормализовать денежное значение до 2 знаков"
            },
        ],
    })

    # 23. SimpleList_Order
    sections.append({
        "kind": "Класс",
        "name": "SimpleList_Order",
        "desc": "Упрощённый динамический контейнер для хранения Order.",
        "fields": [
            {"code": "Order* data_", "desc": "буфер элементов"},
            {"code": "size_t size_", "desc": "текущее количество элементов"},
            {"code": "size_t capacity_", "desc": "текущая ёмкость буфера"},
        ],
        "methods": [
            {"code": "SimpleList_Order()", "desc": "конструктор по умолчанию"},
            {"code": "~SimpleList_Order()", "desc": "деструктор, освобождающий память"},
            {"code": "void push_back(const Order& value)", "desc": "добавить элемент с возможным ростом ёмкости"},
            {"code": "size_t size() const", "desc": "получить количество элементов"},
            {"code": "Order& operator[](size_t idx)", "desc": "доступ к элементу по индексу (неconst)"},
            {"code": "const Order& operator[](size_t idx) const", "desc": "доступ к элементу по индексу (const)"},
            {"code": "Order* begin()", "desc": "итератор на начало"},
            {"code": "Order* end()", "desc": "итератор на конец"},
            {"code": "const Order* begin() const", "desc": "итератор на начало (const)"},
            {"code": "const Order* end() const", "desc": "итератор на конец (const)"},
            {"code": "void clear()", "desc": "очистить контейнер"},
        ],
    })

    # 24. UtilsQt
    sections.append({
        "kind": "Класс",
        "name": "UtilsQt",
        "desc": "Вспомогательные функции для интеграции std::string и Qt и создания кнопок.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические утилиты"},
        ],
        "methods": [
            {"code": "static QString qs(const std::string& s)", "desc": "преобразовать std::string в QString"},
            {"code": "static std::string ss(const QString& s)", "desc": "преобразовать QString в std::string"},
            {
                "code": "static std::string formatName(const std::string& name)",
                "desc": "нормализовать регистр имени (первая буква заглавная)"
            },
            {"code": "static double parsePrice(const QString& text)", "desc": "распарсить строку цены"},
            {
                "code": "static QPushButton* createEditButton(QWidget* parent, const QString& tooltip)",
                "desc": "создать кнопку редактирования с подсказкой"
            },
            {
                "code": "static QPushButton* createDeleteButton(QWidget* parent, const QString& tooltip)",
                "desc": "создать кнопку удаления с подсказкой"
            },
            {
                "code": "static QWidget* createButtonContainer(QWidget* parent, QPushButton* button, bool stretch)",
                "desc": "создать контейнер для кнопки в ячейке таблицы"
            },
        ],
    })

    # 25. AddOrderDialog
    sections.append({
        "kind": "Класс",
        "name": "AddOrderDialog",
        "desc": "Диалог добавления заказа: ввод имени клиента и создание нового заказа.",
        "fields": [
            {"code": "OrderService& svc_", "desc": "сервис заказов"},
            {"code": "QLineEdit* clientEdit_", "desc": "поле ввода имени клиента"},
            {"code": "QDialogButtonBox* buttons_", "desc": "кнопки диалога"},
            {"code": "QCompleter* clientCompleter_", "desc": "автодополнение по клиентам"},
            {"code": "int createdId_", "desc": "идентификатор созданного заказа"},
        ],
        "methods": [
            {
                "code": "AddOrderDialog(OrderService& svc, QWidget* parent)",
                "desc": "конструктор диалога"
            },
            {
                "code": "int createdOrderId() const",
                "desc": "получить идентификатор созданного заказа"
            },
            {
                "code": "void onAdd()",
                "desc": "слот обработки нажатия на кнопку добавления"
            },
        ],
    })

    # 26. AddProductDialog
    sections.append({
        "kind": "Класс",
        "name": "AddProductDialog",
        "desc": "Диалог добавления нового товара.",
        "fields": [
            {"code": "ProductService& productSvc_", "desc": "сервис товаров"},
            {"code": "QLineEdit* nameEdit_", "desc": "поле ввода названия товара"},
            {"code": "QLineEdit* priceEdit_", "desc": "поле ввода цены"},
            {"code": "QLineEdit* stockEdit_", "desc": "поле ввода остатка"},
            {"code": "QDialogButtonBox* buttons_", "desc": "кнопки диалога"},
            {"code": "std::string addedProductName_", "desc": "имя успешно добавленного товара"},
        ],
        "methods": [
            {
                "code": "AddProductDialog(ProductService& productSvc, QWidget* parent)",
                "desc": "конструктор диалога"
            },
            {"code": "void onAdd()", "desc": "слот добавления товара"},
            {
                "code": "std::string addedProductName() const",
                "desc": "получить имя добавленного товара"
            },
        ],
    })

    # 27. EditOrderDialog
    sections.append({
        "kind": "Класс",
        "name": "EditOrderDialog",
        "desc": "Диалог редактирования заказа: статус и позиции.",
        "fields": [
            {"code": "OrderService& svc_", "desc": "сервис заказов"},
            {"code": "ProductService* productSvc_", "desc": "сервис товаров (может быть nullptr)"},
            {"code": "int orderId_", "desc": "идентификатор редактируемого заказа"},
            {"code": "QLineEdit* idEdit_", "desc": "поле отображения ID заказа"},
            {"code": "QComboBox* statusCombo_", "desc": "выбор статуса заказа"},
            {"code": "QTableWidget* itemsTable_", "desc": "таблица позиций заказа"},
            {"code": "QLineEdit* addItemName_", "desc": "поле ввода названия товара для добавления"},
            {"code": "QLineEdit* addQty_", "desc": "поле ввода количества для добавления"},
            {"code": "QPushButton* addItemBtn_", "desc": "кнопка добавления позиции"},
            {"code": "QCompleter* addItemCompleter_", "desc": "автодополнение по товарам"},
        ],
        "methods": [
            {
                "code": "EditOrderDialog(OrderService& svc, int orderId, QWidget* parent)",
                "desc": "конструктор диалога"
            },
            {
                "code": "void setProductService(ProductService* productSvc)",
                "desc": "подключить сервис товаров"
            },
            {"code": "Order* orderOrWarn()", "desc": "получить заказ или показать предупреждение"},
            {"code": "void refreshItemsTable()", "desc": "обновить таблицу позиций"},
            {
                "code": "void onEditItem(const std::string& itemKey, int currentQty)",
                "desc": "редактировать количество позиции"
            },
            {
                "code": "void onDeleteItem(const std::string& itemKey)",
                "desc": "удалить позицию из заказа"
            },
            {
                "code": "void onApplyStatus()",
                "desc": "применить новый статус к заказу"
            },
            {
                "code": "void onAddItem()",
                "desc": "добавить новую позицию в заказ"
            },
        ],
    })

    # 28. FilterDialog
    sections.append({
        "kind": "Класс",
        "name": "FilterDialog",
        "desc": "Диалог настройки фильтров для заказов.",
        "fields": [
            {"code": "QLineEdit* clientEdit_", "desc": "поле ввода фильтра по клиенту"},
            {"code": "QComboBox* statusCombo_", "desc": "комбобокс статуса"},
            {"code": "QLineEdit* minTotalEdit_", "desc": "нижняя граница суммы"},
            {"code": "QLineEdit* maxTotalEdit_", "desc": "верхняя граница суммы"},
            {"code": "QLineEdit* minIdEdit_", "desc": "нижняя граница ID"},
            {"code": "QLineEdit* maxIdEdit_", "desc": "верхняя граница ID"},
            {"code": "QDateTimeEdit* fromDateEdit_", "desc": "нижняя граница даты"},
            {"code": "QDateTimeEdit* toDateEdit_", "desc": "верхняя граница даты"},
            {"code": "QCheckBox* useFromCheck_", "desc": "флаг использования нижней даты"},
            {"code": "QCheckBox* useToCheck_", "desc": "флаг использования верхней даты"},
            {"code": "QDialogButtonBox* buttons_", "desc": "кнопки диалога"},
        ],
        "methods": [
            {
                "code": "FilterDialog(const FilterParams& params, QWidget* parent)",
                "desc": "конструктор с предзаполнением параметров"
            },
            {"code": "QString clientFilter() const", "desc": "получить строку фильтра по клиенту"},
            {"code": "QString statusFilter() const", "desc": "получить строку фильтра по статусу"},
            {"code": "QString minTotalText() const", "desc": "получить текст нижней границы суммы"},
            {"code": "QString maxTotalText() const", "desc": "получить текст верхней границы суммы"},
            {"code": "QString minIdText() const", "desc": "получить текст нижней границы ID"},
            {"code": "QString maxIdText() const", "desc": "получить текст верхней границы ID"},
            {"code": "bool useFrom() const", "desc": "используется ли нижняя граница даты"},
            {"code": "bool useTo() const", "desc": "используется ли верхняя граница даты"},
            {"code": "QDateTime fromDate() const", "desc": "получить нижнюю дату"},
            {"code": "QDateTime toDate() const", "desc": "получить верхнюю дату"},
        ],
    })

    # 29. ReportDialog
    sections.append({
        "kind": "Класс",
        "name": "ReportDialog",
        "desc": "Диалог выбора параметров текстового отчёта.",
        "fields": [
            {"code": "QLineEdit* nameEdit_", "desc": "поле ввода имени отчёта"},
            {"code": "QComboBox* scopeCombo_", "desc": "выбор области отчёта (фильтр/всё)"},
            {"code": "QCheckBox* includeFilters_", "desc": "флаг включения раздела с фильтрами"},
            {"code": "QCheckBox* includeSummary_", "desc": "флаг включения сводной информации"},
        ],
        "methods": [
            {
                "code": "ReportDialog(bool filterActive, QWidget* parent)",
                "desc": "конструктор, учитывающий активность фильтрации"
            },
            {
                "code": "QString reportName() const",
                "desc": "получить имя отчёта"
            },
            {
                "code": "bool scopeFiltered() const",
                "desc": "должен ли отчёт формироваться только по отфильтрованным данным"
            },
            {
                "code": "bool includeFiltersHeader() const",
                "desc": "включать ли раздел с параметрами фильтра"
            },
            {
                "code": "bool includeSummarySection() const",
                "desc": "включать ли сводный раздел"
            },
        ],
    })

    # 30. MainWindow
    sections.append({
        "kind": "Класс",
        "name": "MainWindow",
        "desc": "Главное окно приложения: таблица заказов и товаров, фильтры, отчёты и статистика.",
        "fields": [
            {"code": "OrderService& svc_", "desc": "сервис заказов"},
            {"code": "ProductService& productSvc_", "desc": "сервис товаров"},
            {"code": "ValidationService V_", "desc": "сервис валидации"},
            {"code": "QTabWidget* tabs_", "desc": "вкладки интерфейса"},
            {"code": "QTableWidget* table_", "desc": "таблица заказов"},
            {"code": "QTableWidget* productTable_", "desc": "таблица товаров"},
            {"code": "QPushButton* addOrderBtn_", "desc": "кнопка добавления заказа"},
            {"code": "QPushButton* reportBtn_", "desc": "кнопка формирования отчёта"},
            {"code": "QPushButton* addProductBtn_", "desc": "кнопка добавления товара"},
            {"code": "QPushButton* clearFilterBtn_", "desc": "кнопка сброса фильтров"},
            {"code": "QPushButton* openChartsBtn_", "desc": "кнопка открытия статистики"},
            {"code": "QLabel* titleLabel_", "desc": "заголовок окна"},
            {"code": "StatisticsWindow* statisticsWindow_", "desc": "окно статистики"},
            {"code": "QCompleter* clientFilterCompleter_", "desc": "автодополнение по клиентам"},
            {"code": "QStringListModel* clientFilterModel_", "desc": "модель клиентов для фильтрации"},
            {"code": "OrderStatsLabels orderStats_", "desc": "лейблы статистики по заказам"},
            {"code": "ProductStatsLabels productStats_", "desc": "лейблы статистики по товарам"},
            {"code": "FilterWidgets filterWidgets_", "desc": "указатели на виджеты фильтра"},
            {"code": "FilterState filterState_", "desc": "текущее состояние фильтрации"},
        ],
        "methods": [
            {
                "code": "MainWindow(OrderService& svc, ProductService& productSvc, QWidget* parent)",
                "desc": "конструктор главного окна"
            },
            {"code": "void refreshTable()", "desc": "обновить таблицу заказов"},
            {"code": "void refreshProducts()", "desc": "обновить таблицу товаров"},
            {
                "code": "QList<const Order*> currentFilteredRows() const",
                "desc": "получить список заказов, прошедших фильтр"
            },
            {"code": "void applyFilters()", "desc": "применить фильтры к таблице заказов"},
            {"code": "void setupCompleters()", "desc": "настроить автодополнения"},
            {
                "code": "void onAddOrder()",
                "desc": "обработчик добавления заказа"
            },
            {
                "code": "void onOpenReportDialog()",
                "desc": "открыть диалог формирования отчёта"
            },
            {
                "code": "void onOpenStatistics()",
                "desc": "открыть окно статистики"
            },
            {
                "code": "void onClearFilter()",
                "desc": "сбросить фильтры"
            },
            {
                "code": "void onFilterChanged()",
                "desc": "обработать изменение фильтров"
            },
            {
                "code": "void updateStatistics()",
                "desc": "обновить сводную статистику по заказам"
            },
            {
                "code": "void updateProductStatistics()",
                "desc": "обновить статистику по товарам"
            },
            {
                "code": "void onAddProduct()",
                "desc": "обработать добавление товара"
            },
            {
                "code": "void resizeEvent(QResizeEvent* event)",
                "desc": "реакция на изменение размера окна"
            },
        ],
    })

    # 31. StatisticsWindow
    sections.append({
        "kind": "Класс",
        "name": "StatisticsWindow",
        "desc": "Окно статистики по заказам с анимацией и диаграммами.",
        "fields": [
            {"code": "OrderService& svc_", "desc": "сервис заказов"},
            {"code": "QTimer* animationTimer_", "desc": "таймер анимации"},
            {"code": "double animationProgress_", "desc": "прогресс анимации"},
        ],
        "methods": [
            {
                "code": "StatisticsWindow(OrderService& svc, QWidget* parent)",
                "desc": "конструктор окна статистики"
            },
            {"code": "void refreshStatistics()", "desc": "пересчитать статистику и перерисовать окно"},
            {"code": "void showEvent(QShowEvent* event)", "desc": "обработать показ окна"},
            {"code": "void paintEvent(QPaintEvent* event)", "desc": "перерисовать диаграммы"},
            {
                "code": "void onAnimationTick()",
                "desc": "шаг анимации (обновление прогресса)"
            },
            {
                "code": "void drawBarWithGradient(QPainter& painter, const QRect& rect, const QColor& color, bool selected)",
                "desc": "отрисовать один столбец статистики с градиентом"
            },
            {
                "code": "void drawAnimatedRevenueBar(QPainter& painter, const RevenueBarParams& params)",
                "desc": "отрисовать анимированный столбец выручки"
            },
        ],
    })

    # 32. ProductWindow
    sections.append({
        "kind": "Класс",
        "name": "ProductWindow",
        "desc": "Окно управления товарами: просмотр, добавление и редактирование.",
        "fields": [
            {"code": "ProductService& productSvc_", "desc": "сервис товаров"},
            {"code": "OrderService& orderSvc_", "desc": "сервис заказов"},
            {"code": "QTableWidget* productTable_", "desc": "таблица товаров"},
            {"code": "QPushButton* addProductBtn_", "desc": "кнопка добавления товара"},
            {"code": "QCompleter* productNameCompleter_", "desc": "автодополнение по названию товара"},
        ],
        "methods": [
            {
                "code": "ProductWindow(ProductService& productSvc, OrderService& orderSvc, QWidget* parent)",
                "desc": "конструктор окна товаров"
            },
            {"code": "void refreshProducts()", "desc": "обновить список товаров"},
            {"code": "void setupCompleters()", "desc": "настроить автодополнения"},
            {"code": "void onAddProduct()", "desc": "обработчик добавления товара"},
            {
                "code": "void onEditProduct(std::string_view productKey, std::string_view productName)",
                "desc": "обработчик редактирования товара"
            },
            {
                "code": "void onDeleteProduct(const std::string& productKey, const std::string& productName)",
                "desc": "обработчик удаления товара"
            },
            {
                "code": "bool isProductUsedInActiveOrders(const std::string& productKey, QList<int>& affectedOrderIds) const",
                "desc": "проверить, используется ли товар в активных заказах"
            },
            {"code": "void showEvent(QShowEvent* event)", "desc": "обработчик показа окна"},
        ],
    })

    # 33. NumericItem
    sections.append({
        "kind": "Класс",
        "name": "NumericItem",
        "desc": "Элемент таблицы Qt с корректной числовой сортировкой.",
        "fields": [
            {"code": "double v_", "desc": "числовое значение для сравнения"},
        ],
        "methods": [
            {
                "code": "NumericItem(int v, const QString& display)",
                "desc": "конструктор из целого значения"
            },
            {
                "code": "NumericItem(double v, const QString& display)",
                "desc": "конструктор из дробного значения"
            },
            {
                "code": "std::partial_ordering operator<=>(const QTableWidgetItem& other) const",
                "desc": "сравнение с другим элементом по числовому значению"
            },
        ],
    })

    return sections


def build_docx(output_path: str):
    doc = Document()

    # Заголовок раздела
    make_header_paragraph(doc, "3.1 Описание программных модулей")

    sections = build_sections()
    for idx, sec in enumerate(sections, start=1):
        make_class_heading(doc, idx, sec["kind"], sec["name"])
        make_normal_paragraph(doc, sec["desc"])
        add_bullet_block(doc, "Поля:", sec["fields"])
        add_bullet_block(doc, "Методы:", sec["methods"])

    doc.save(output_path)


if __name__ == "__main__":
    build_docx("Раздел_3_1_классы_и_структуры.docx")
    print("Готово: Раздел_3_1_классы_и_структуры.docx")

