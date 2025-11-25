from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

NBSP = "\u00A0"  # неразрывный пробел


def set_font(run, name: str, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(14)
    run.font.bold = bold


def make_header_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_normal_paragraph(doc: Document, text: str):
    """Обычный абзац с красной строкой 1.25 см, TNR 14."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_class_heading(doc: Document, number: int, kind: str, name: str):
    """
    Заголовок: "N Класс Name:" или "N Структура Name:"
    N + 'Класс/Структура' — Times New Roman
    Name: — Courier New
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)

    r1 = p.add_run(f"{number} {kind} ")
    set_font(r1, "Times New Roman")

    r2 = p.add_run(f"{name}:")
    set_font(r2, "Courier New")

    return p


def add_bullet_block(doc: Document, title: str, items: list[dict]):
    """
    items: список словарей вида {"code": "...", "desc": "..."}.
    code — шрифт кода (Courier New), desc — Times New Roman.
    """
    # Заголовок "Поля:" / "Методы:"
    make_normal_paragraph(doc, title)

    if not items:
        return

    n = len(items)
    for idx, item in enumerate(items):
        last = (idx == n - 1)
        p = doc.add_paragraph()
        pf = p.paragraph_format

        # ВАЖНО: убираем красную строку принудительно
        pf.left_indent = Cm(2.5)
        pf.first_line_indent = Cm(0)   # ← отключаем красную строку
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)

        code = item.get("code", "").strip()
        desc = (item.get("desc", "") or "").strip()

        # Маркер "– " (с NBSP), Times New Roman
        r0 = p.add_run("–" + NBSP)
        set_font(r0, "Times New Roman")

        # Код — Courier New
        if code:
            r_code = p.add_run(code)
            set_font(r_code, "Courier New")

        # Описание — Times New Roman
        ending = "." if last else ";"

        if desc:
            r_colon = p.add_run(": ")
            set_font(r_colon, "Times New Roman")

            r_desc = p.add_run(desc + ending)
            set_font(r_desc, "Times New Roman")
        else:
            # Нет описания — просто код + точка/точка с запятой
            r_end = p.add_run(ending)
            set_font(r_end, "Times New Roman")


def build_sections() -> list[dict]:
    """
    Возвращает список всех сущностей (классы/структуры) с полями/методами.
    Формат элемента:
    {
        "kind": "Класс" или "Структура",
        "name": "Имя",
        "desc": "Описание",
        "fields": [{"code": "...", "desc": "..."}, ...],
        "methods": [{"code": "...", "desc": "..."}, ...],
    }
    """
    sections: list[dict] = []

    # 1. Property
    sections.append({
        "kind": "Класс",
        "name": "Property",
        "desc": "Базовый абстрактный класс для объектов недвижимости: идентификатор, адрес, цена, площадь, описание и доступность.",
        "fields": [
            {"code": "std::string id", "desc": "идентификатор объекта недвижимости"},
            {"code": "std::string city", "desc": "город"},
            {"code": "std::string street", "desc": "улица"},
            {"code": "std::string house", "desc": "номер дома"},
            {"code": "double price", "desc": "цена объекта"},
            {"code": "double area", "desc": "площадь объекта"},
            {"code": "std::string description", "desc": "описание объекта"},
            {"code": "bool isAvailable", "desc": "флаг доступности объекта"},
        ],
        "methods": [
            {
                "code": "Property(const std::string& id, const std::string& city, const std::string& street, const std::string& house, double price, double area, const std::string& description)",
                "desc": "конструктор базового класса"
            },
            {"code": "virtual ~Property()", "desc": "виртуальный деструктор"},
            {"code": "virtual std::string getType() const = 0", "desc": "получить тип недвижимости (чисто виртуальный)"},
            {"code": "virtual void displayInfo() const = 0", "desc": "вывести информацию об объекте (чисто виртуальный)"},
            {"code": "virtual std::string toFileString() const = 0", "desc": "сериализация в строку для файла (чисто виртуальный)"},
            {"code": "virtual Property* clone() const = 0", "desc": "клонирование объекта (чисто виртуальный)"},
            {"code": "bool operator==(const Property& other) const", "desc": "проверка равенства по идентификатору"},
            {"code": "std::partial_ordering operator<=>(const Property& other) const", "desc": "сравнение объектов"},
            {"code": "std::string getId() const", "desc": "получить идентификатор"},
            {"code": "std::string getCity() const", "desc": "получить город"},
            {"code": "std::string getStreet() const", "desc": "получить улицу"},
            {"code": "std::string getHouse() const", "desc": "получить номер дома"},
            {"code": "std::string getAddress() const", "desc": "получить полный адрес"},
            {"code": "double getPrice() const", "desc": "получить цену"},
            {"code": "double getArea() const", "desc": "получить площадь"},
            {"code": "std::string getDescription() const", "desc": "получить описание"},
            {"code": "bool getIsAvailable() const", "desc": "проверить доступность"},
            {"code": "void setPrice(double newPrice)", "desc": "установить цену"},
            {"code": "void setArea(double newArea)", "desc": "установить площадь"},
            {"code": "void setAddress(const std::string& city, const std::string& street, const std::string& house)", "desc": "установить адрес"},
            {"code": "void setDescription(std::string_view newDesc)", "desc": "установить описание"},
            {"code": "void setAvailable(bool available)", "desc": "установить доступность"},
            {"code": "static bool validatePrice(double price)", "desc": "валидация цены"},
            {"code": "static bool validateArea(double area)", "desc": "валидация площади"},
            {"code": "static bool validateId(std::string_view id)", "desc": "валидация идентификатора"},
            {"code": "static bool validateAddressPart(std::string_view part)", "desc": "валидация части адреса"},
            {"code": "friend std::ostream& operator<<(std::ostream& os, const Property& prop)", "desc": "вывод объекта в поток"},
        ],
    })

    # 2. Apartment
    sections.append({
        "kind": "Класс",
        "name": "Apartment",
        "desc": "Класс квартиры, наследуется от Property: количество комнат, этаж, наличие балкона и лифта.",
        "fields": [
            {"code": "int rooms", "desc": "количество комнат"},
            {"code": "int floor", "desc": "этаж"},
            {"code": "bool hasBalcony", "desc": "наличие балкона"},
            {"code": "bool hasElevator", "desc": "наличие лифта"},
        ],
        "methods": [
            {"code": "explicit Apartment(const ApartmentParams& params)", "desc": "конструктор из параметров"},
            {"code": "std::string getType() const override", "desc": "получить тип (\"Apartment\")"},
            {"code": "void displayInfo() const override", "desc": "вывести информацию о квартире"},
            {"code": "std::string toFileString() const override", "desc": "сериализация в строку для файла"},
            {"code": "Property* clone() const override", "desc": "клонирование квартиры"},
            {"code": "int getRooms() const", "desc": "получить количество комнат"},
            {"code": "int getFloor() const", "desc": "получить этаж"},
            {"code": "bool getHasBalcony() const", "desc": "проверить наличие балкона"},
            {"code": "bool getHasElevator() const", "desc": "проверить наличие лифта"},
            {"code": "void setRooms(int rooms)", "desc": "установить количество комнат"},
            {"code": "void setFloor(int floor)", "desc": "установить этаж"},
        ],
    })

    # 3. House
    sections.append({
        "kind": "Класс",
        "name": "House",
        "desc": "Класс дома, наследуется от Property: количество этажей, комнат, площадь участка, наличие гаража и сада.",
        "fields": [
            {"code": "int floors", "desc": "количество этажей"},
            {"code": "int rooms", "desc": "количество комнат"},
            {"code": "double landArea", "desc": "площадь участка"},
            {"code": "bool hasGarage", "desc": "наличие гаража"},
            {"code": "bool hasGarden", "desc": "наличие сада"},
        ],
        "methods": [
            {"code": "explicit House(const HouseParams& params)", "desc": "конструктор из параметров"},
            {"code": "std::string getType() const override", "desc": "получить тип (\"House\")"},
            {"code": "void displayInfo() const override", "desc": "вывести информацию о доме"},
            {"code": "std::string toFileString() const override", "desc": "сериализация в строку для файла"},
            {"code": "Property* clone() const override", "desc": "клонирование дома"},
            {"code": "int getFloors() const", "desc": "получить количество этажей"},
            {"code": "int getRooms() const", "desc": "получить количество комнат"},
            {"code": "double getLandArea() const", "desc": "получить площадь участка"},
            {"code": "bool getHasGarage() const", "desc": "проверить наличие гаража"},
            {"code": "bool getHasGarden() const", "desc": "проверить наличие сада"},
            {"code": "void setFloors(int floors)", "desc": "установить количество этажей"},
            {"code": "void setRooms(int rooms)", "desc": "установить количество комнат"},
            {"code": "void setLandArea(double landArea)", "desc": "установить площадь участка"},
        ],
    })

    # 4. CommercialProperty
    sections.append({
        "kind": "Класс",
        "name": "CommercialProperty",
        "desc": "Класс коммерческой недвижимости, наследуется от Property: тип бизнеса, наличие парковки, количество парковочных мест, видимость с улицы.",
        "fields": [
            {"code": "std::string businessType", "desc": "тип бизнеса"},
            {"code": "bool hasParking", "desc": "наличие парковки"},
            {"code": "int parkingSpaces", "desc": "количество парковочных мест"},
            {"code": "bool isVisibleFromStreet", "desc": "видимость с улицы"},
        ],
        "methods": [
            {"code": "explicit CommercialProperty(const CommercialPropertyParams& params)", "desc": "конструктор из параметров"},
            {"code": "std::string getType() const override", "desc": "получить тип (\"Commercial\")"},
            {"code": "void displayInfo() const override", "desc": "вывести информацию о коммерческой недвижимости"},
            {"code": "std::string toFileString() const override", "desc": "сериализация в строку для файла"},
            {"code": "Property* clone() const override", "desc": "клонирование коммерческой недвижимости"},
            {"code": "std::string getBusinessType() const", "desc": "получить тип бизнеса"},
            {"code": "bool getHasParking() const", "desc": "проверить наличие парковки"},
            {"code": "int getParkingSpaces() const", "desc": "получить количество парковочных мест"},
            {"code": "bool getIsVisibleFromStreet() const", "desc": "проверить видимость с улицы"},
            {"code": "void setBusinessType(std::string_view type)", "desc": "установить тип бизнеса"},
            {"code": "void setParkingSpaces(int spaces)", "desc": "установить количество парковочных мест"},
        ],
    })

    # 5. Client
    sections.append({
        "kind": "Класс",
        "name": "Client",
        "desc": "Класс клиента: идентификатор, имя, телефон, email и дата регистрации.",
        "fields": [
            {"code": "std::string id", "desc": "идентификатор клиента"},
            {"code": "std::string name", "desc": "имя клиента"},
            {"code": "std::string phone", "desc": "телефон клиента"},
            {"code": "std::string email", "desc": "email клиента"},
            {"code": "std::string registrationDate", "desc": "дата регистрации в формате ISO 8601"},
        ],
        "methods": [
            {"code": "Client(const std::string& id, const std::string& name, const std::string& phone, const std::string& email)", "desc": "конструктор клиента"},
            {"code": "bool operator==(const Client& other) const", "desc": "проверка равенства по идентификатору"},
            {"code": "std::strong_ordering operator<=>(const Client& other) const", "desc": "сравнение клиентов"},
            {"code": "std::string getId() const", "desc": "получить идентификатор"},
            {"code": "std::string getName() const", "desc": "получить имя"},
            {"code": "std::string getPhone() const", "desc": "получить телефон"},
            {"code": "std::string getEmail() const", "desc": "получить email"},
            {"code": "std::string getRegistrationDate() const", "desc": "получить дату регистрации"},
            {"code": "void setName(std::string_view name)", "desc": "установить имя"},
            {"code": "void setPhone(std::string_view phone)", "desc": "установить телефон"},
            {"code": "void setEmail(std::string_view email)", "desc": "установить email"},
            {"code": "static bool validateId(std::string_view id)", "desc": "валидация идентификатора"},
            {"code": "static bool validatePhone(std::string_view phone)", "desc": "валидация телефона"},
            {"code": "static bool validateEmail(std::string_view email)", "desc": "валидация email"},
            {"code": "std::string toString() const", "desc": "преобразование в строку"},
            {"code": "friend std::ostream& operator<<(std::ostream& os, const Client& client)", "desc": "вывод клиента в поток"},
        ],
    })

    # 6. Transaction
    sections.append({
        "kind": "Класс",
        "name": "Transaction",
        "desc": "Класс сделки: идентификатор, идентификатор объекта недвижимости, идентификатор клиента, дата, финальная цена, статус и примечания.",
        "fields": [
            {"code": "std::string id", "desc": "идентификатор сделки"},
            {"code": "std::string propertyId", "desc": "идентификатор объекта недвижимости"},
            {"code": "std::string clientId", "desc": "идентификатор клиента"},
            {"code": "std::string date", "desc": "дата сделки в формате ISO 8601"},
            {"code": "double finalPrice", "desc": "финальная цена сделки"},
            {"code": "std::string status", "desc": "статус сделки (pending, completed, cancelled)"},
            {"code": "std::string notes", "desc": "примечания к сделке"},
        ],
        "methods": [
            {"code": "Transaction(const std::string& id, const std::string& propertyId, const std::string& clientId, double finalPrice, const std::string& status = \"pending\", const std::string& notes = \"\")", "desc": "конструктор сделки"},
            {"code": "bool operator==(const Transaction& other) const", "desc": "проверка равенства по идентификатору"},
            {"code": "std::strong_ordering operator<=>(const Transaction& other) const", "desc": "сравнение сделок"},
            {"code": "std::string getId() const", "desc": "получить идентификатор"},
            {"code": "std::string getPropertyId() const", "desc": "получить идентификатор объекта недвижимости"},
            {"code": "std::string getClientId() const", "desc": "получить идентификатор клиента"},
            {"code": "std::string getDate() const", "desc": "получить дату сделки"},
            {"code": "double getFinalPrice() const", "desc": "получить финальную цену"},
            {"code": "std::string getStatus() const", "desc": "получить статус"},
            {"code": "std::string getNotes() const", "desc": "получить примечания"},
            {"code": "void setStatus(std::string_view status)", "desc": "установить статус"},
            {"code": "void setFinalPrice(double price)", "desc": "установить финальную цену"},
            {"code": "void setNotes(std::string_view notes)", "desc": "установить примечания"},
            {"code": "static bool validateId(std::string_view id)", "desc": "валидация идентификатора"},
            {"code": "std::string toString() const", "desc": "преобразование в строку"},
            {"code": "std::string toFileString() const", "desc": "сериализация в строку для файла"},
            {"code": "friend std::ostream& operator<<(std::ostream& os, const Transaction& trans)", "desc": "вывод сделки в поток"},
        ],
    })

    # 7. Auction
    sections.append({
        "kind": "Класс",
        "name": "Auction",
        "desc": "Класс аукциона: идентификатор, идентификатор объекта недвижимости, адрес объекта, начальная цена, цена выкупа, список ставок, статус и даты создания и завершения.",
        "fields": [
            {"code": "std::string id", "desc": "идентификатор аукциона"},
            {"code": "std::string propertyId", "desc": "идентификатор объекта недвижимости"},
            {"code": "std::string propertyAddress", "desc": "адрес объекта недвижимости"},
            {"code": "double startingPrice", "desc": "начальная цена"},
            {"code": "double buyoutPrice", "desc": "цена выкупа"},
            {"code": "std::vector<std::shared_ptr<Bid>> bids", "desc": "список ставок"},
            {"code": "std::string status", "desc": "статус аукциона (active, completed, cancelled)"},
            {"code": "std::string createdAt", "desc": "дата создания в формате ISO 8601"},
            {"code": "std::string completedAt", "desc": "дата завершения в формате ISO 8601"},
        ],
        "methods": [
            {"code": "Auction(const std::string& id, const std::string& propertyId, const std::string& propertyAddress, double startingPrice)", "desc": "конструктор аукциона"},
            {"code": "bool operator==(const Auction& other) const", "desc": "проверка равенства по идентификатору"},
            {"code": "std::strong_ordering operator<=>(const Auction& other) const", "desc": "сравнение аукционов"},
            {"code": "bool addBid(std::shared_ptr<Bid> bid)", "desc": "добавить ставку с проверкой"},
            {"code": "void addBidDirect(std::shared_ptr<Bid> bid)", "desc": "добавить ставку напрямую"},
            {"code": "double getCurrentHighestBid() const", "desc": "получить текущую максимальную ставку"},
            {"code": "const Bid* getHighestBid() const", "desc": "получить указатель на максимальную ставку"},
            {"code": "void complete()", "desc": "завершить аукцион"},
            {"code": "void cancel()", "desc": "отменить аукцион"},
            {"code": "std::string getId() const", "desc": "получить идентификатор"},
            {"code": "std::string getPropertyId() const", "desc": "получить идентификатор объекта недвижимости"},
            {"code": "std::string getPropertyAddress() const", "desc": "получить адрес объекта"},
            {"code": "double getStartingPrice() const", "desc": "получить начальную цену"},
            {"code": "double getBuyoutPrice() const", "desc": "получить цену выкупа"},
            {"code": "std::vector<std::shared_ptr<Bid>> getBids() const", "desc": "получить список ставок"},
            {"code": "std::string getStatus() const", "desc": "получить статус"},
            {"code": "std::string getCreatedAt() const", "desc": "получить дату создания"},
            {"code": "std::string getCompletedAt() const", "desc": "получить дату завершения"},
            {"code": "bool isActive() const", "desc": "проверить, активен ли аукцион"},
            {"code": "bool isCompleted() const", "desc": "проверить, завершен ли аукцион"},
            {"code": "bool wasBuyout() const", "desc": "проверить, был ли выкуп"},
            {"code": "std::string toString() const", "desc": "преобразование в строку"},
            {"code": "std::string toFileString() const", "desc": "сериализация в строку для файла"},
            {"code": "friend std::ostream& operator<<(std::ostream& os, const Auction& auction)", "desc": "вывод аукциона в поток"},
        ],
    })

    # 8. Bid
    sections.append({
        "kind": "Класс",
        "name": "Bid",
        "desc": "Класс ставки в аукционе: идентификатор клиента, имя клиента, сумма ставки и временная метка.",
        "fields": [
            {"code": "std::string clientId", "desc": "идентификатор клиента"},
            {"code": "std::string clientName", "desc": "имя клиента"},
            {"code": "double amount", "desc": "сумма ставки"},
            {"code": "std::string timestamp", "desc": "временная метка в формате ISO 8601"},
        ],
        "methods": [
            {"code": "Bid(const std::string& clientId, const std::string& clientName, double amount)", "desc": "конструктор ставки"},
            {"code": "bool operator==(const Bid& other) const", "desc": "проверка равенства ставок"},
            {"code": "std::partial_ordering operator<=>(const Bid& other) const", "desc": "сравнение ставок"},
            {"code": "std::string getClientId() const", "desc": "получить идентификатор клиента"},
            {"code": "std::string getClientName() const", "desc": "получить имя клиента"},
            {"code": "double getAmount() const", "desc": "получить сумму ставки"},
            {"code": "std::string getTimestamp() const", "desc": "получить временную метку"},
            {"code": "std::string toString() const", "desc": "преобразование в строку"},
            {"code": "std::string toFileString() const", "desc": "сериализация в строку для файла"},
            {"code": "friend std::ostream& operator<<(std::ostream& os, const Bid& bid)", "desc": "вывод ставки в поток"},
        ],
    })

    # 9. PropertyBaseParams
    sections.append({
        "kind": "Структура",
        "name": "PropertyBaseParams",
        "desc": "Базовые параметры объекта недвижимости для создания.",
        "fields": [
            {"code": "std::string id", "desc": "идентификатор объекта"},
            {"code": "std::string city", "desc": "город"},
            {"code": "std::string street", "desc": "улица"},
            {"code": "std::string house", "desc": "номер дома"},
            {"code": "double price", "desc": "цена"},
            {"code": "double area", "desc": "площадь"},
            {"code": "std::string description", "desc": "описание"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения параметров"},
        ],
    })

    # 10. ApartmentParams
    sections.append({
        "kind": "Структура",
        "name": "ApartmentParams",
        "desc": "Параметры квартиры для создания, включает базовые параметры и специфичные для квартиры.",
        "fields": [
            {"code": "PropertyBaseParams base", "desc": "базовые параметры объекта недвижимости"},
            {"code": "int rooms", "desc": "количество комнат"},
            {"code": "int floor", "desc": "этаж"},
            {"code": "bool hasBalcony", "desc": "наличие балкона"},
            {"code": "bool hasElevator", "desc": "наличие лифта"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения параметров"},
        ],
    })

    # 11. HouseParams
    sections.append({
        "kind": "Структура",
        "name": "HouseParams",
        "desc": "Параметры дома для создания, включает базовые параметры и специфичные для дома.",
        "fields": [
            {"code": "PropertyBaseParams base", "desc": "базовые параметры объекта недвижимости"},
            {"code": "int floors", "desc": "количество этажей"},
            {"code": "int rooms", "desc": "количество комнат"},
            {"code": "double landArea", "desc": "площадь участка"},
            {"code": "bool hasGarage", "desc": "наличие гаража"},
            {"code": "bool hasGarden", "desc": "наличие сада"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения параметров"},
        ],
    })

    # 12. CommercialPropertyParams
    sections.append({
        "kind": "Структура",
        "name": "CommercialPropertyParams",
        "desc": "Параметры коммерческой недвижимости для создания, включает базовые параметры и специфичные для коммерческой недвижимости.",
        "fields": [
            {"code": "PropertyBaseParams base", "desc": "базовые параметры объекта недвижимости"},
            {"code": "std::string businessType", "desc": "тип бизнеса"},
            {"code": "bool hasParking", "desc": "наличие парковки"},
            {"code": "int parkingSpaces", "desc": "количество парковочных мест"},
            {"code": "bool isVisibleFromStreet", "desc": "видимость с улицы"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения параметров"},
        ],
    })

    # 13. PropertyManager
    sections.append({
        "kind": "Класс",
        "name": "PropertyManager",
        "desc": "Менеджер объектов недвижимости: управление коллекцией объектов, добавление, удаление, поиск и фильтрация.",
        "fields": [
            {"code": "std::vector<std::unique_ptr<Property>> properties", "desc": "коллекция объектов недвижимости"},
        ],
        "methods": [
            {"code": "PropertyManager()", "desc": "конструктор менеджера"},
            {"code": "void addProperty(std::unique_ptr<Property> property)", "desc": "добавить объект недвижимости"},
            {"code": "void addApartment(const ApartmentParams& params)", "desc": "добавить квартиру"},
            {"code": "void addHouse(const HouseParams& params)", "desc": "добавить дом"},
            {"code": "void addCommercialProperty(const CommercialPropertyParams& params)", "desc": "добавить коммерческую недвижимость"},
            {"code": "bool removeProperty(const std::string& id)", "desc": "удалить объект недвижимости по идентификатору"},
            {"code": "Property* findProperty(const std::string& id) const", "desc": "найти объект по идентификатору"},
            {"code": "std::vector<Property*> getAllProperties() const", "desc": "получить все объекты недвижимости"},
            {"code": "std::vector<Property*> getAvailableProperties() const", "desc": "получить доступные объекты недвижимости"},
            {"code": "std::vector<Property*> searchByPriceRange(double minPrice, double maxPrice) const", "desc": "поиск по диапазону цен"},
            {"code": "std::vector<Property*> searchByAddress(const std::string& city, const std::string& street = \"\", const std::string& house = \"\") const", "desc": "поиск по адресу"},
            {"code": "const std::vector<std::unique_ptr<Property>>& getProperties() const", "desc": "получить коллекцию объектов для FileManager"},
            {"code": "void setProperties(std::vector<std::unique_ptr<Property>>&& props)", "desc": "установить коллекцию объектов из FileManager"},
            {"code": "size_t getCount() const", "desc": "получить количество объектов"},
        ],
    })

    # 14. ClientManager
    sections.append({
        "kind": "Класс",
        "name": "ClientManager",
        "desc": "Менеджер клиентов: управление коллекцией клиентов, добавление, удаление и поиск.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<Client>> clients", "desc": "коллекция клиентов"},
        ],
        "methods": [
            {"code": "ClientManager()", "desc": "конструктор менеджера"},
            {"code": "void addClient(std::shared_ptr<Client> client)", "desc": "добавить клиента"},
            {"code": "bool removeClient(const std::string& id)", "desc": "удалить клиента по идентификатору"},
            {"code": "Client* findClient(const std::string& id) const", "desc": "найти клиента по идентификатору"},
            {"code": "std::vector<Client*> getAllClients() const", "desc": "получить всех клиентов"},
            {"code": "std::vector<Client*> searchByName(std::string_view name) const", "desc": "поиск по имени"},
            {"code": "std::vector<Client*> searchByPhone(std::string_view phone) const", "desc": "поиск по телефону"},
            {"code": "const std::vector<std::shared_ptr<Client>>& getClients() const", "desc": "получить коллекцию клиентов для FileManager"},
            {"code": "void setClients(std::vector<std::shared_ptr<Client>>&& newClients)", "desc": "установить коллекцию клиентов из FileManager"},
            {"code": "size_t getCount() const", "desc": "получить количество клиентов"},
        ],
    })

    # 15. TransactionManager
    sections.append({
        "kind": "Класс",
        "name": "TransactionManager",
        "desc": "Менеджер сделок: управление коллекцией сделок, добавление, удаление и поиск по различным критериям.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<Transaction>> transactions", "desc": "коллекция сделок"},
        ],
        "methods": [
            {"code": "TransactionManager()", "desc": "конструктор менеджера"},
            {"code": "void addTransaction(std::shared_ptr<Transaction> transaction)", "desc": "добавить сделку"},
            {"code": "bool removeTransaction(const std::string& id)", "desc": "удалить сделку по идентификатору"},
            {"code": "Transaction* findTransaction(const std::string& id) const", "desc": "найти сделку по идентификатору"},
            {"code": "std::vector<Transaction*> getAllTransactions() const", "desc": "получить все сделки"},
            {"code": "std::vector<Transaction*> getTransactionsByClient(std::string_view clientId) const", "desc": "получить сделки клиента"},
            {"code": "std::vector<Transaction*> getTransactionsByProperty(std::string_view propertyId) const", "desc": "получить сделки по объекту недвижимости"},
            {"code": "std::vector<Transaction*> getTransactionsByStatus(std::string_view status) const", "desc": "получить сделки по статусу"},
            {"code": "const std::vector<std::shared_ptr<Transaction>>& getTransactions() const", "desc": "получить коллекцию сделок для FileManager"},
            {"code": "void setTransactions(std::vector<std::shared_ptr<Transaction>>&& newTransactions)", "desc": "установить коллекцию сделок из FileManager"},
            {"code": "size_t getCount() const", "desc": "получить количество сделок"},
        ],
    })

    # 16. AuctionManager
    sections.append({
        "kind": "Класс",
        "name": "AuctionManager",
        "desc": "Менеджер аукционов: управление коллекцией аукционов, добавление, удаление и поиск.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<Auction>> auctions", "desc": "коллекция аукционов"},
        ],
        "methods": [
            {"code": "AuctionManager()", "desc": "конструктор менеджера"},
            {"code": "void addAuction(std::shared_ptr<Auction> auction)", "desc": "добавить аукцион"},
            {"code": "bool removeAuction(const std::string& id)", "desc": "удалить аукцион по идентификатору"},
            {"code": "Auction* findAuction(const std::string& id) const", "desc": "найти аукцион по идентификатору"},
            {"code": "std::vector<Auction*> getAllAuctions() const", "desc": "получить все аукционы"},
            {"code": "std::vector<Auction*> getActiveAuctions() const", "desc": "получить активные аукционы"},
            {"code": "std::vector<Auction*> getCompletedAuctions() const", "desc": "получить завершенные аукционы"},
            {"code": "std::vector<Auction*> getAuctionsByProperty(std::string_view propertyId) const", "desc": "получить аукционы по объекту недвижимости"},
            {"code": "const std::vector<std::shared_ptr<Auction>>& getAuctions() const", "desc": "получить коллекцию аукционов для FileManager"},
            {"code": "void setAuctions(std::vector<std::shared_ptr<Auction>>&& newAuctions)", "desc": "установить коллекцию аукционов из FileManager"},
            {"code": "size_t getCount() const", "desc": "получить количество аукционов"},
        ],
    })

    # 17. FileManager
    sections.append({
        "kind": "Класс",
        "name": "FileManager",
        "desc": "Менеджер файлов: сохранение и загрузка данных в текстовый формат с разметкой через менеджеры.",
        "fields": [
            {"code": "static constexpr char FILE_DELIMITER", "desc": "разделитель полей в файле"},
            {"code": "static constexpr char AVAILABLE_CHAR", "desc": "символ доступности"},
            {"code": "static constexpr char UNAVAILABLE_CHAR", "desc": "символ недоступности"},
            {"code": "static constexpr size_t BID_PREFIX_LENGTH", "desc": "длина префикса ставки"},
        ],
        "methods": [
            {"code": "static void saveProperties(const PropertyManager& manager, const std::string& filename)", "desc": "сохранить объекты недвижимости в файл"},
            {"code": "static void saveClients(const ClientManager& manager, const std::string& filename)", "desc": "сохранить клиентов в файл"},
            {"code": "static void saveTransactions(const TransactionManager& manager, const std::string& filename)", "desc": "сохранить сделки в файл"},
            {"code": "static void saveAuctions(const AuctionManager& manager, const std::string& filename)", "desc": "сохранить аукционы в файл"},
            {"code": "static void loadProperties(PropertyManager& manager, const std::string& filename)", "desc": "загрузить объекты недвижимости из файла"},
            {"code": "static void loadClients(ClientManager& manager, const std::string& filename)", "desc": "загрузить клиентов из файла"},
            {"code": "static void loadTransactions(TransactionManager& manager, const std::string& filename)", "desc": "загрузить сделки из файла"},
            {"code": "static void loadAuctions(AuctionManager& manager, const std::string& filename)", "desc": "загрузить аукционы из файла"},
            {"code": "static void parseBidLine(std::string_view line, Auction* currentAuction)", "desc": "парсинг строки ставки из файла"},
            {"code": "static std::shared_ptr<Auction> parseAuctionLine(std::string_view line)", "desc": "парсинг строки аукциона из файла"},
        ],
    })

    # 18. EstateAgency
    sections.append({
        "kind": "Класс",
        "name": "EstateAgency",
        "desc": "Центральный класс-синглтон агентства недвижимости: управление всеми менеджерами и координация работы системы.",
        "fields": [
            {"code": "static EstateAgency* instance", "desc": "единственный экземпляр синглтона"},
            {"code": "PropertyManager propertyManager", "desc": "менеджер объектов недвижимости"},
            {"code": "ClientManager clientManager", "desc": "менеджер клиентов"},
            {"code": "TransactionManager transactionManager", "desc": "менеджер сделок"},
            {"code": "AuctionManager auctionManager", "desc": "менеджер аукционов"},
            {"code": "std::string dataDirectory", "desc": "директория для хранения данных"},
        ],
        "methods": [
            {"code": "EstateAgency()", "desc": "приватный конструктор синглтона"},
            {"code": "~EstateAgency()", "desc": "деструктор синглтона"},
            {"code": "static EstateAgency* getInstance()", "desc": "получить экземпляр синглтона"},
            {"code": "static void destroyInstance()", "desc": "уничтожить экземпляр синглтона"},
            {"code": "PropertyManager& getPropertyManager()", "desc": "получить менеджер объектов недвижимости"},
            {"code": "ClientManager& getClientManager()", "desc": "получить менеджер клиентов"},
            {"code": "TransactionManager& getTransactionManager()", "desc": "получить менеджер сделок"},
            {"code": "AuctionManager& getAuctionManager()", "desc": "получить менеджер аукционов"},
            {"code": "void saveAllData() const", "desc": "сохранить все данные через FileManager"},
            {"code": "void loadAllData()", "desc": "загрузить все данные через FileManager"},
            {"code": "void setDataDirectory(std::string_view dir)", "desc": "установить директорию для данных"},
            {"code": "std::string getDataDirectory() const", "desc": "получить директорию для данных"},
        ],
    })

    # 19. PropertyManagerException
    sections.append({
        "kind": "Класс",
        "name": "PropertyManagerException",
        "desc": "Исключение для ошибок PropertyManager, наследуется от std::exception.",
        "fields": [
            {"code": "std::string message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit PropertyManagerException(const std::string& msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
        ],
    })

    # 20. ClientManagerException
    sections.append({
        "kind": "Класс",
        "name": "ClientManagerException",
        "desc": "Исключение для ошибок ClientManager, наследуется от std::exception.",
        "fields": [
            {"code": "std::string message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit ClientManagerException(const std::string& msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
        ],
    })

    # 21. TransactionManagerException
    sections.append({
        "kind": "Класс",
        "name": "TransactionManagerException",
        "desc": "Исключение для ошибок TransactionManager, наследуется от std::exception.",
        "fields": [
            {"code": "std::string message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit TransactionManagerException(const std::string& msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
        ],
    })

    # 22. AuctionManagerException
    sections.append({
        "kind": "Класс",
        "name": "AuctionManagerException",
        "desc": "Исключение для ошибок AuctionManager, наследуется от std::exception.",
        "fields": [
            {"code": "std::string message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit AuctionManagerException(const std::string& msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
        ],
    })

    # 23. FileManagerException
    sections.append({
        "kind": "Класс",
        "name": "FileManagerException",
        "desc": "Исключение для ошибок FileManager, наследуется от std::exception.",
        "fields": [
            {"code": "std::string message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit FileManagerException(const std::string& msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
        ],
    })

    # 24. ContainerManager
    sections.append({
        "kind": "Класс",
        "name": "ContainerManager",
        "desc": "Шаблонный класс-контейнер для управления коллекциями объектов с идентификаторами.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<T>> items", "desc": "коллекция элементов"},
        ],
        "methods": [
            {"code": "void add(std::shared_ptr<T> item)", "desc": "добавить элемент"},
            {"code": "void remove(const std::string& id)", "desc": "удалить элемент по идентификатору"},
            {"code": "std::shared_ptr<T> findById(const std::string& id) const", "desc": "найти элемент по идентификатору"},
            {"code": "size_t size() const", "desc": "получить количество элементов"},
            {"code": "typename std::vector<std::shared_ptr<T>>::const_iterator begin() const", "desc": "итератор на начало"},
            {"code": "typename std::vector<std::shared_ptr<T>>::const_iterator end() const", "desc": "итератор на конец"},
        ],
    })

    # 25. MainWindow
    sections.append({
        "kind": "Класс",
        "name": "MainWindow",
        "desc": "Главное окно приложения: навигация, виджеты для работы с данными, обработка исключений и сохранение/загрузка данных.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на экземпляр EstateAgency"},
            {"code": "QListWidget* navigationList", "desc": "виджет навигации"},
            {"code": "QStackedWidget* contentStack", "desc": "стек виджетов для переключения разделов"},
            {"code": "DashboardWidget* dashboardWidget", "desc": "виджет панели управления"},
            {"code": "PropertiesWidget* propertiesWidget", "desc": "виджет управления объектами недвижимости"},
            {"code": "ClientsWidget* clientsWidget", "desc": "виджет управления клиентами"},
            {"code": "TransactionsWidget* transactionsWidget", "desc": "виджет управления сделками"},
            {"code": "AuctionsWidget* auctionsWidget", "desc": "виджет управления аукционами"},
        ],
        "methods": [
            {"code": "explicit MainWindow(QWidget* parent = nullptr)", "desc": "конструктор главного окна"},
            {"code": "~MainWindow() override", "desc": "деструктор главного окна"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void setupMenuBar()", "desc": "настройка меню"},
            {"code": "void setupNewUI()", "desc": "настройка нового интерфейса"},
            {"code": "void applyStyles()", "desc": "применение стилей"},
            {"code": "void updateDashboardStats()", "desc": "обновление статистики на панели"},
            {"code": "void showStatusMessage(const QString& message, int timeout = 3000) const", "desc": "показать сообщение в статусной строке"},
            {"code": "void handleException(const std::exception& e)", "desc": "обработка исключений"},
            {"code": "template<typename Func> void executeWithExceptionHandling(Func&& operation)", "desc": "выполнение операции с обработкой исключений"},
            {"code": "void saveAllData()", "desc": "сохранить все данные"},
            {"code": "void loadAllData()", "desc": "загрузить все данные"},
            {"code": "void refreshAllData()", "desc": "обновить все данные"},
            {"code": "void onNavigationChanged(int index)", "desc": "обработчик изменения навигации"},
            {"code": "void onDataChanged()", "desc": "обработчик изменения данных"},
        ],
    })

    # 26. PropertyDialog
    sections.append({
        "kind": "Класс",
        "name": "PropertyDialog",
        "desc": "Диалог добавления и редактирования объекта недвижимости: форма с общими полями и динамическими полями в зависимости от типа.",
        "fields": [
            {"code": "enum class PropertyType", "desc": "тип недвижимости (TypeApartment, TypeHouse, TypeCommercial)"},
            {"code": "struct CommonFields", "desc": "общие поля для всех типов недвижимости"},
            {"code": "struct ApartmentFields", "desc": "поля специфичные для квартиры"},
            {"code": "struct HouseFields", "desc": "поля специфичные для дома"},
            {"code": "struct CommercialFields", "desc": "поля специфичные для коммерческой недвижимости"},
            {"code": "QDialogButtonBox* buttonBox", "desc": "кнопки диалога"},
        ],
        "methods": [
            {"code": "explicit PropertyDialog(QWidget* parent = nullptr, const Property* editProperty = nullptr)", "desc": "конструктор диалога"},
            {"code": "~PropertyDialog() override", "desc": "деструктор диалога"},
            {"code": "PropertyType getPropertyType() const", "desc": "получить тип недвижимости"},
            {"code": "QString getId() const", "desc": "получить идентификатор"},
            {"code": "QString getCity() const", "desc": "получить город"},
            {"code": "QString getStreet() const", "desc": "получить улицу"},
            {"code": "QString getHouse() const", "desc": "получить номер дома"},
            {"code": "double getPrice() const", "desc": "получить цену"},
            {"code": "double getArea() const", "desc": "получить площадь"},
            {"code": "QString getDescription() const", "desc": "получить описание"},
            {"code": "bool getIsAvailable() const", "desc": "проверить доступность"},
            {"code": "int getRooms() const", "desc": "получить количество комнат (квартира)"},
            {"code": "int getFloor() const", "desc": "получить этаж (квартира)"},
            {"code": "bool getHasBalcony() const", "desc": "проверить наличие балкона (квартира)"},
            {"code": "bool getHasElevator() const", "desc": "проверить наличие лифта (квартира)"},
            {"code": "int getFloors() const", "desc": "получить количество этажей (дом)"},
            {"code": "double getLandArea() const", "desc": "получить площадь участка (дом)"},
            {"code": "bool getHasGarage() const", "desc": "проверить наличие гаража (дом)"},
            {"code": "bool getHasGarden() const", "desc": "проверить наличие сада (дом)"},
            {"code": "QString getBusinessType() const", "desc": "получить тип бизнеса (коммерческая)"},
            {"code": "int getParkingSpaces() const", "desc": "получить количество парковочных мест (коммерческая)"},
            {"code": "bool getIsVisibleFromStreet() const", "desc": "проверить видимость с улицы (коммерческая)"},
            {"code": "bool getHasParking() const", "desc": "проверить наличие парковки (коммерческая)"},
            {"code": "void propertyTypeChanged(int index)", "desc": "обработчик изменения типа недвижимости"},
            {"code": "void validateAndAccept()", "desc": "валидация и принятие данных"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void loadPropertyData(const Property* prop)", "desc": "загрузка данных объекта для редактирования"},
            {"code": "void updateTypeSpecificFields()", "desc": "обновление полей в зависимости от типа"},
        ],
    })

    # 27. ClientDialog
    sections.append({
        "kind": "Класс",
        "name": "ClientDialog",
        "desc": "Диалог добавления и редактирования клиента: форма с полями идентификатора, имени, телефона и email.",
        "fields": [
            {"code": "QLineEdit* idEdit", "desc": "поле ввода идентификатора"},
            {"code": "QLineEdit* nameEdit", "desc": "поле ввода имени"},
            {"code": "QLineEdit* phoneEdit", "desc": "поле ввода телефона"},
            {"code": "QLineEdit* emailEdit", "desc": "поле ввода email"},
            {"code": "QDialogButtonBox* buttonBox", "desc": "кнопки диалога"},
        ],
        "methods": [
            {"code": "explicit ClientDialog(QWidget* parent = nullptr, const Client* editClient = nullptr)", "desc": "конструктор диалога"},
            {"code": "~ClientDialog() override", "desc": "деструктор диалога"},
            {"code": "QString getId() const", "desc": "получить идентификатор"},
            {"code": "QString getName() const", "desc": "получить имя"},
            {"code": "QString getPhone() const", "desc": "получить телефон"},
            {"code": "QString getEmail() const", "desc": "получить email"},
            {"code": "void validateAndAccept()", "desc": "валидация и принятие данных"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void loadClientData(const Client* client)", "desc": "загрузка данных клиента для редактирования"},
        ],
    })

    # 28. TransactionDialog
    sections.append({
        "kind": "Класс",
        "name": "TransactionDialog",
        "desc": "Диалог добавления и редактирования сделки: форма с выбором объекта недвижимости и клиента, ценой, статусом и примечаниями.",
        "fields": [
            {"code": "QLineEdit* idEdit", "desc": "поле ввода идентификатора сделки"},
            {"code": "QComboBox* propertyCombo", "desc": "выбор объекта недвижимости"},
            {"code": "QComboBox* clientCombo", "desc": "выбор клиента"},
            {"code": "QDoubleSpinBox* priceSpin", "desc": "поле ввода финальной цены"},
            {"code": "QComboBox* statusCombo", "desc": "выбор статуса сделки"},
            {"code": "QTextEdit* notesEdit", "desc": "поле ввода примечаний"},
            {"code": "QDialogButtonBox* buttonBox", "desc": "кнопки диалога"},
            {"code": "QLabel* propertyPriceLabel", "desc": "метка с ценой объекта"},
            {"code": "QLabel* differenceLabel", "desc": "метка с разницей цен"},
            {"code": "QStringList propertyIds", "desc": "список идентификаторов объектов"},
            {"code": "QStringList clientIds", "desc": "список идентификаторов клиентов"},
            {"code": "bool isUpdatingFromProperty", "desc": "флаг обновления из объекта"},
        ],
        "methods": [
            {"code": "explicit TransactionDialog(QWidget* parent = nullptr, const Transaction* editTransaction = nullptr, const QStringList& propertyIds = QStringList(), const QStringList& clientIds = QStringList())", "desc": "конструктор диалога"},
            {"code": "~TransactionDialog() override", "desc": "деструктор диалога"},
            {"code": "QString getId() const", "desc": "получить идентификатор"},
            {"code": "QString getPropertyId() const", "desc": "получить идентификатор объекта недвижимости"},
            {"code": "QString getClientId() const", "desc": "получить идентификатор клиента"},
            {"code": "double getFinalPrice() const", "desc": "получить финальную цену"},
            {"code": "QString getStatus() const", "desc": "получить статус"},
            {"code": "QString getNotes() const", "desc": "получить примечания"},
            {"code": "void validateAndAccept()", "desc": "валидация и принятие данных"},
            {"code": "void onPropertyChanged(int index)", "desc": "обработчик изменения объекта недвижимости"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void loadTransactionData(const Transaction* trans)", "desc": "загрузка данных сделки для редактирования"},
            {"code": "void updatePriceFromProperty()", "desc": "обновление цены из объекта недвижимости"},
            {"code": "void updatePriceDifference()", "desc": "обновление разницы цен"},
        ],
    })

    # 29. AuctionDialog
    sections.append({
        "kind": "Класс",
        "name": "AuctionDialog",
        "desc": "Диалог создания и просмотра аукциона: форма с выбором объекта недвижимости, начальной ценой, таблицей ставок и управлением аукционом.",
        "fields": [
            {"code": "bool isViewMode", "desc": "режим просмотра (только чтение)"},
            {"code": "Auction* currentAuction", "desc": "текущий редактируемый аукцион"},
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QLineEdit* idEdit", "desc": "поле ввода идентификатора"},
            {"code": "QComboBox* propertyCombo", "desc": "выбор объекта недвижимости"},
            {"code": "QDoubleSpinBox* priceSpin", "desc": "поле ввода начальной цены"},
            {"code": "QLabel* propertyPriceLabel", "desc": "метка с ценой объекта"},
            {"code": "QLabel* buyoutPriceLabel", "desc": "метка с ценой выкупа"},
            {"code": "QLabel* currentHighestBidLabel", "desc": "метка с текущей максимальной ставкой"},
            {"code": "QLabel* statusLabel", "desc": "метка со статусом"},
            {"code": "QTableWidget* bidsTable", "desc": "таблица ставок"},
            {"code": "QPushButton* addBidBtn", "desc": "кнопка добавления ставки"},
            {"code": "QPushButton* completeAuctionBtn", "desc": "кнопка завершения аукциона"},
            {"code": "QComboBox* clientCombo", "desc": "выбор клиента для ставки"},
            {"code": "QDoubleSpinBox* bidAmountSpin", "desc": "поле ввода суммы ставки"},
            {"code": "QDialogButtonBox* buttonBox", "desc": "кнопки диалога"},
            {"code": "QStringList propertyIds", "desc": "список идентификаторов объектов"},
        ],
        "methods": [
            {"code": "explicit AuctionDialog(QWidget* parent = nullptr, Auction* editAuction = nullptr, const QStringList& propertyIds = QStringList())", "desc": "конструктор диалога"},
            {"code": "~AuctionDialog() override", "desc": "деструктор диалога"},
            {"code": "QString getId() const", "desc": "получить идентификатор"},
            {"code": "QString getPropertyId() const", "desc": "получить идентификатор объекта недвижимости"},
            {"code": "double getStartingPrice() const", "desc": "получить начальную цену"},
            {"code": "void refreshBids()", "desc": "обновить таблицу ставок"},
            {"code": "void updateAuctionInfo()", "desc": "обновить информацию об аукционе"},
            {"code": "void validateAndAccept()", "desc": "валидация и принятие данных"},
            {"code": "void onPropertyChanged(int index)", "desc": "обработчик изменения объекта недвижимости"},
            {"code": "void addBid()", "desc": "добавить ставку"},
            {"code": "void completeAuction()", "desc": "завершить аукцион"},
            {"code": "void refreshAuctionInfo()", "desc": "обновить информацию об аукционе"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void loadAuctionData(const Auction* auction)", "desc": "загрузка данных аукциона для просмотра"},
            {"code": "void updatePropertyInfo()", "desc": "обновление информации об объекте недвижимости"},
            {"code": "void createTransactionFromAuction()", "desc": "создание сделки из завершенного аукциона"},
        ],
    })

    # 30. PropertiesWidget
    sections.append({
        "kind": "Класс",
        "name": "PropertiesWidget",
        "desc": "Виджет управления объектами недвижимости: таблица объектов, добавление, редактирование, удаление, поиск и отображение деталей.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QTableWidget* propertiesTable", "desc": "таблица объектов недвижимости"},
            {"code": "QPushButton* addPropertyBtn", "desc": "кнопка добавления объекта"},
            {"code": "QPushButton* refreshPropertyBtn", "desc": "кнопка обновления"},
            {"code": "QPushButton* searchPropertyBtn", "desc": "кнопка поиска"},
            {"code": "QLineEdit* searchPropertyEdit", "desc": "поле ввода поискового запроса"},
            {"code": "QTextEdit* propertyDetailsText", "desc": "текстовое поле для деталей объекта"},
        ],
        "methods": [
            {"code": "explicit PropertiesWidget(EstateAgency* agency, QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void refresh()", "desc": "обновить данные"},
            {"code": "void updateTable()", "desc": "обновить таблицу"},
            {"code": "void addProperty()", "desc": "добавить объект недвижимости"},
            {"code": "void editProperty()", "desc": "редактировать объект недвижимости"},
            {"code": "void deleteProperty()", "desc": "удалить объект недвижимости"},
            {"code": "void searchProperties()", "desc": "поиск объектов недвижимости"},
            {"code": "void propertySelectionChanged()", "desc": "обработчик изменения выбора в таблице"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void showPropertyDetails(const Property* prop)", "desc": "показать детали объекта"},
            {"code": "void showPropertyTransactions(const std::string& propertyId)", "desc": "показать сделки по объекту"},
            {"code": "void addPropertyToTable(const Property* prop)", "desc": "добавить объект в таблицу"},
            {"code": "void selectRowById(QTableWidget* table, const QString& id) const", "desc": "выбрать строку по идентификатору"},
            {"code": "QString getSelectedIdFromTable(const QTableWidget* table) const", "desc": "получить идентификатор выбранной строки"},
            {"code": "bool checkTableSelection(const QTableWidget* table, const QString& errorMessage)", "desc": "проверить выбор в таблице"},
            {"code": "bool isNumericId(const QString& text) const", "desc": "проверить, является ли текст числовым идентификатором"},
        ],
    })

    # 31. ClientsWidget
    sections.append({
        "kind": "Класс",
        "name": "ClientsWidget",
        "desc": "Виджет управления клиентами: таблица клиентов, добавление, редактирование, удаление, поиск и отображение деталей.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QTableWidget* clientsTable", "desc": "таблица клиентов"},
            {"code": "QPushButton* addClientBtn", "desc": "кнопка добавления клиента"},
            {"code": "QPushButton* refreshClientBtn", "desc": "кнопка обновления"},
            {"code": "QPushButton* searchClientBtn", "desc": "кнопка поиска"},
            {"code": "QLineEdit* searchClientEdit", "desc": "поле ввода поискового запроса"},
            {"code": "QTextEdit* clientDetailsText", "desc": "текстовое поле для деталей клиента"},
        ],
        "methods": [
            {"code": "explicit ClientsWidget(EstateAgency* agency, QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void refresh()", "desc": "обновить данные"},
            {"code": "void updateTable()", "desc": "обновить таблицу"},
            {"code": "void addClient()", "desc": "добавить клиента"},
            {"code": "void editClient()", "desc": "редактировать клиента"},
            {"code": "void deleteClient()", "desc": "удалить клиента"},
            {"code": "void searchClients()", "desc": "поиск клиентов"},
            {"code": "void clientSelectionChanged()", "desc": "обработчик изменения выбора в таблице"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void showClientDetails(const Client* client)", "desc": "показать детали клиента"},
            {"code": "void showClientTransactions(const std::string& clientId)", "desc": "показать сделки клиента"},
            {"code": "void addClientToTable(const Client* client)", "desc": "добавить клиента в таблицу"},
            {"code": "void selectRowById(QTableWidget* table, const QString& id) const", "desc": "выбрать строку по идентификатору"},
            {"code": "QString getSelectedIdFromTable(const QTableWidget* table) const", "desc": "получить идентификатор выбранной строки"},
            {"code": "bool checkTableSelection(const QTableWidget* table, const QString& errorMessage)", "desc": "проверить выбор в таблице"},
        ],
    })

    # 32. TransactionsWidget
    sections.append({
        "kind": "Класс",
        "name": "TransactionsWidget",
        "desc": "Виджет управления сделками: таблица сделок, добавление, редактирование, удаление, поиск и отображение деталей.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QTableWidget* transactionsTable", "desc": "таблица сделок"},
            {"code": "QPushButton* addTransactionBtn", "desc": "кнопка добавления сделки"},
            {"code": "QPushButton* refreshTransactionBtn", "desc": "кнопка обновления"},
            {"code": "QPushButton* searchTransactionBtn", "desc": "кнопка поиска"},
            {"code": "QLineEdit* searchTransactionEdit", "desc": "поле ввода поискового запроса"},
            {"code": "QTextEdit* transactionDetailsText", "desc": "текстовое поле для деталей сделки"},
        ],
        "methods": [
            {"code": "explicit TransactionsWidget(EstateAgency* agency, QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void refresh()", "desc": "обновить данные"},
            {"code": "void updateTable()", "desc": "обновить таблицу"},
            {"code": "void addTransaction()", "desc": "добавить сделку"},
            {"code": "void editTransaction()", "desc": "редактировать сделку"},
            {"code": "void deleteTransaction()", "desc": "удалить сделку"},
            {"code": "void searchTransactions()", "desc": "поиск сделок"},
            {"code": "void transactionSelectionChanged()", "desc": "обработчик изменения выбора в таблице"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void showTransactionDetails(const Transaction* trans)", "desc": "показать детали сделки"},
            {"code": "void addTransactionToTable(const Transaction* trans)", "desc": "добавить сделку в таблицу"},
            {"code": "bool validateTransaction(std::string_view propertyId, std::string_view clientId, std::string_view status, std::string_view excludeId = \"\")", "desc": "валидация сделки"},
            {"code": "bool hasActiveTransactions(const std::string& propertyId)", "desc": "проверить наличие активных сделок по объекту"},
            {"code": "void selectRowById(QTableWidget* table, const QString& id) const", "desc": "выбрать строку по идентификатору"},
            {"code": "QString getSelectedIdFromTable(const QTableWidget* table) const", "desc": "получить идентификатор выбранной строки"},
            {"code": "bool checkTableSelection(const QTableWidget* table, const QString& errorMessage)", "desc": "проверить выбор в таблице"},
        ],
    })

    # 33. AuctionsWidget
    sections.append({
        "kind": "Класс",
        "name": "AuctionsWidget",
        "desc": "Виджет управления аукционами: таблица аукционов, создание, просмотр, удаление, поиск и отображение деталей.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QTableWidget* auctionsTable", "desc": "таблица аукционов"},
            {"code": "QPushButton* addAuctionBtn", "desc": "кнопка создания аукциона"},
            {"code": "QPushButton* refreshAuctionBtn", "desc": "кнопка обновления"},
            {"code": "QPushButton* searchAuctionBtn", "desc": "кнопка поиска"},
            {"code": "QLineEdit* searchAuctionEdit", "desc": "поле ввода поискового запроса"},
            {"code": "QTextEdit* auctionDetailsText", "desc": "текстовое поле для деталей аукциона"},
        ],
        "methods": [
            {"code": "explicit AuctionsWidget(EstateAgency* agency, QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void refresh()", "desc": "обновить данные"},
            {"code": "void updateTable()", "desc": "обновить таблицу"},
            {"code": "void addAuction()", "desc": "создать аукцион"},
            {"code": "void viewAuction()", "desc": "просмотреть аукцион"},
            {"code": "void deleteAuction()", "desc": "удалить аукцион"},
            {"code": "void searchAuctions()", "desc": "поиск аукционов"},
            {"code": "void auctionSelectionChanged()", "desc": "обработчик изменения выбора в таблице"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
            {"code": "void showAuctionDetails(const Auction* auction)", "desc": "показать детали аукциона"},
            {"code": "void addAuctionToTable(const Auction* auction)", "desc": "добавить аукцион в таблицу"},
            {"code": "bool hasActiveTransactions(const std::string& propertyId)", "desc": "проверить наличие активных сделок по объекту"},
            {"code": "QWidget* createActionButtons(QTableWidget* table, const QString& id, const std::function<void()>& viewAction, const std::function<void()>& deleteAction, bool isView = false)", "desc": "создать кнопки действий в таблице"},
            {"code": "void selectRowById(QTableWidget* table, const QString& id) const", "desc": "выбрать строку по идентификатору"},
            {"code": "QString getSelectedIdFromTable(const QTableWidget* table) const", "desc": "получить идентификатор выбранной строки"},
            {"code": "bool checkTableSelection(const QTableWidget* table, const QString& errorMessage)", "desc": "проверить выбор в таблице"},
        ],
    })

    # 34. DashboardWidget
    sections.append({
        "kind": "Класс",
        "name": "DashboardWidget",
        "desc": "Виджет панели управления: отображение статистики по объектам недвижимости, клиентам, сделкам и аукционам.",
        "fields": [
            {"code": "EstateAgency* agency", "desc": "указатель на EstateAgency"},
            {"code": "QLabel* statsPropertiesLabel", "desc": "метка со статистикой по объектам недвижимости"},
            {"code": "QLabel* statsClientsLabel", "desc": "метка со статистикой по клиентам"},
            {"code": "QLabel* statsTransactionsLabel", "desc": "метка со статистикой по сделкам"},
            {"code": "QLabel* statsAvailableLabel", "desc": "метка со статистикой по доступным объектам"},
            {"code": "QLabel* statsAuctionsLabel", "desc": "метка со статистикой по аукционам"},
            {"code": "QPushButton* saveBtn", "desc": "кнопка сохранения данных"},
            {"code": "QPushButton* loadBtn", "desc": "кнопка загрузки данных"},
            {"code": "QPushButton* refreshBtn", "desc": "кнопка обновления статистики"},
        ],
        "methods": [
            {"code": "explicit DashboardWidget(EstateAgency* agency, QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void updateStats()", "desc": "обновить статистику"},
            {"code": "void setupUI()", "desc": "настройка интерфейса"},
        ],
    })

    return sections


def build_docx(output_path: str):
    doc = Document()

    # Заголовок раздела
    make_header_paragraph(doc, "3.1 Описание программных модулей")

    sections = build_sections()
    for idx, sec in enumerate(sections, start=1):
        make_class_heading(doc, idx, sec["kind"], sec["name"])
        make_normal_paragraph(doc, sec["desc"])
        add_bullet_block(doc, "Поля:", sec["fields"])
        add_bullet_block(doc, "Методы:", sec["methods"])

    doc.save(output_path)


if __name__ == "__main__":
    build_docx("Раздел_3_1_классы_и_структуры.docx")
    print("Готово: Раздел_3_1_классы_и_структуры.docx")


